"""Build the formal configuration and operations report from its Markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
MUTED = "666666"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
BLACK = "000000"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=None, bold=None, color=None, italic=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    code = styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_pr = code._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F6F8FA")
    p_pr.append(shd)

    quote = styles.add_style("Report Quote", 1)
    quote.font.name = "Calibri"
    quote.font.size = Pt(10.5)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor.from_string(MUTED)
    quote._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.15)
    quote.paragraph_format.space_after = Pt(8)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TABsucks | 软件配置与运维文档")
    set_font(run, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("内部项目文档  |  第 ")
    set_font(run, size=8.5, color=MUTED)
    add_field(p, "PAGE", "1")
    run = p.add_run(" 页，共 ")
    set_font(run, size=8.5, color=MUTED)
    add_field(p, "NUMPAGES", "1")
    run = p.add_run(" 页")
    set_font(run, size=8.5, color=MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("TABsucks")
    set_font(run, size=17, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("软件配置与运维文档")
    set_font(run, size=28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(58)
    run = p.add_run("Configuration Management and Operations Plan")
    set_font(run, size=12, color=MUTED, italic=True)

    rows = [
        ("文档编号", "TABSUCKS-CMO-001"),
        ("文档版本", "V1.0"),
        ("文档状态", "课程提交候选版"),
        ("项目团队", "TABsucks 项目组"),
        ("组长学号/姓名", "待填写"),
        ("编制日期", "2026-07-18"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    set_table_geometry(table, [2500, 6860])
    for index, (label, value) in enumerate(rows):
        for cell in table.rows[index].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(index, 0), PALE_BLUE)
        p0 = table.cell(index, 0).paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(label)
        set_font(r0, size=10.5, bold=True, color=NAVY)
        p1 = table.cell(index, 1).paragraphs[0]
        r1 = p1.add_run(value)
        set_font(r1, size=10.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("适用范围：TABsucks Windows 桌面发行与本地 Web 运行环境")
    set_font(run, size=9.5, color=MUTED)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph("目录", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "在 Word 中右键更新目录")
    doc.add_page_break()


INLINE_PATTERN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")


def add_inline_runs(paragraph, text: str, *, size=None, color=None) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=size or 9.5, color="8B1A1A")
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
        else:
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size, bold=True, color=color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, size=size, color=color)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return [rows[0], *rows[2:]]


def table_widths(rows: list[list[str]]) -> list[int]:
    columns = len(rows[0])
    weights = []
    for index in range(columns):
        maximum = max(len(row[index]) if index < len(row) else 0 for row in rows)
        weights.append(max(8, min(maximum, 38)))
    total = sum(weights)
    raw = [max(900, int(CONTENT_WIDTH_DXA * weight / total)) for weight in weights]
    difference = CONTENT_WIDTH_DXA - sum(raw)
    raw[-1] += difference
    return raw


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    widths = table_widths(rows)
    set_table_geometry(table, widths)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if row_index == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_shading(cell, LIGHT_GRAY)
            add_inline_runs(p, value, size=9 if len(rows[0]) >= 4 else 9.5)
            for run in p.runs:
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    first_title_skipped = False
    body_started = False
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip() or line.strip() == "---":
            index += 1
            continue
        if line.startswith("```"):
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            p = doc.add_paragraph(style="Code Block")
            run = p.add_run("\n".join(code_lines))
            set_font(run, size=8.5, east_asia="Microsoft YaHei")
            run.font.name = "Consolas"
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(
            r"^\s*\|?\s*:?-+", lines[index + 1]
        ):
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, parse_table(table_lines))
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not first_title_skipped:
                first_title_skipped = True
                index += 1
                continue
            if not body_started:
                if level == 2 and text == "文档控制":
                    body_started = True
                else:
                    index += 1
                    continue
            p = doc.add_paragraph(style=f"Heading {level}")
            p.paragraph_format.keep_with_next = True
            add_inline_runs(p, text)
            index += 1
            continue
        if not body_started:
            index += 1
            continue
        if line.startswith("> "):
            p = doc.add_paragraph(style="Report Quote")
            add_inline_runs(p, line[2:])
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, bullet.group(1))
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, numbered.group(1))
            index += 1
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            candidate = lines[index].rstrip()
            if (
                not candidate.strip()
                or candidate.startswith(("#", "```", "|", "> ", "- "))
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(part.strip() for part in paragraph_lines))


def build(source: Path, output: Path) -> None:
    markdown = source.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    add_header_footer(doc.sections[0])
    add_cover(doc)
    add_toc(doc)
    add_markdown(doc, markdown)
    configure_sections(doc)
    for section in doc.sections:
        add_header_footer(section)
    doc.core_properties.title = "TABsucks 软件配置与运维文档"
    doc.core_properties.subject = "软件配置管理、版本控制、持续集成、部署与运维"
    doc.core_properties.author = "TABsucks 项目组"
    doc.core_properties.keywords = "TABsucks, 配置管理, 运维, 发布, Windows"
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_configuration_operations_doc.py SOURCE.md OUTPUT.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))


if __name__ == "__main__":
    main()
